from PyQt5 import QtGui
import sys
from PyQt5 import QtCore, QtWidgets
from PyQt5.QtWidgets import QTableWidgetItem, QMessageBox, QFileDialog
from ui import Ui_MainWindow
import MySQLdb as Mdb
import inspect
from string import digits
from docx import Document
from docxtpl import DocxTemplate
import datetime
import openpyxl
from openpyxl.styles import Font
from openpyxl.styles import Border, Side
from register import Ui_Form


class Program:
    def __init__(self):
        self.register_ui = Ui_Form()
        self.form2 = From2()

        self.form2.window.enterEnterButton.clicked.connect(self.enter_button_clicked)

        self.form2_show()

    def enter_button_clicked(self):
        message1 = QMessageBox(MainWindow)
        message1.setIcon(QMessageBox.Warning)
        message1.setWindowTitle("Неправильные данные")

        login = self.form2.window.enterLoginEdit.text()
        password = self.form2.window.enterPassEdit.text()

        enter_data = []
        role = "Пользователь"
        for record in self.form2.data_table:
            if record[1] == login and record[2] == password:
                role = record[3]
                enter_data = [login, password, role]

        if len(enter_data) == 0:
            message1.setText("Не найден логин или пароль")
            message1.exec_()
            return

        employee_id = 0
        for record in self.form2.data_table:
            if record[2] == password:
                employee_id = record[4]

        self.form1_show(login, role, employee_id)

    def form1_show(self, login, role, employee_id):
        self.form2.close()
        self.main_ui = Ui_MainWindow()
        self.form1 = From1(login, role, employee_id)
        self.main_ui.setupUi(MainWindow)
        self.form1.show()

    def form2_show(self):
        self.register_ui.setupUi(MainWindow)
        self.form2.show()


class From2(QtWidgets.QMainWindow):
    def __init__(self):
        super(From2, self).__init__()
        self.window = Ui_Form()
        self.window.setupUi(self)

        self.admin_password = "Хочу 9"

        self.db = Mdb.connect('localhost', 'root', '321Ilyxazc', 'enter_data')
        self.db.set_character_set("utf8")
        self.cursor = self.db.cursor()
        self.cursor.execute('''SELECT * FROM данныеВхода''')
        self.data_table = self.cursor.fetchall()

        self.form_initialization()

    def form_initialization(self):
        self.window.regAdminPassEdit.setVisible(False)
        self.window.label_6.setVisible(False)
        self.window.tabWidget.setCurrentIndex(0)

        trpo_db = Mdb.connect('localhost', 'root', '321Ilyxazc', 'trpo_db')
        trpo_db.set_character_set("utf8")
        cursor = trpo_db.cursor()
        cursor.execute('''SELECT * FROM сотрудники''')
        self.staff_table = cursor.fetchall()

        for employee in self.staff_table:
            self.window.regIdCB.addItem(str(employee[0]))

        self.window.enterPassEdit.setEchoMode(QtWidgets.QLineEdit.Password)
        self.window.regPassFEdit.setEchoMode(QtWidgets.QLineEdit.Password)
        self.window.regPassSEdit.setEchoMode(QtWidgets.QLineEdit.Password)
        self.window.regAdminPassEdit.setEchoMode(QtWidgets.QLineEdit.Password)

        self.window.regRegisterButton.clicked.connect(self.reg_register_button_clicked)
        self.window.regAdminRadioB.toggled.connect(self.reg_admin_radio_button_toggled)
        self.window.enterGoToRegButton.clicked.connect(self.enter_go_to_reg_button_clicked)
        self.window.regReturnToEnterButton.clicked.connect(self.reg_return_to_enter_button_clicked)

    def update_table(self):
        self.cursor.execute('''SELECT * FROM данныеВхода''')
        self.data_table = self.cursor.fetchall()

    def check_in_table(self, data):
        for record in self.data_table:
            if record[1] == data:
                return True
        return False

    def insert_into_table(self, login, password, role, employee_id):
        sql = f'''INSERT INTO `enter_data`.`данныевхода` (`Логин`, `Пароль`, `Роль`, `Код`)
         VALUES ('{login}', '{password}', '{role}', '{employee_id}');'''
        self.cursor.execute(sql)
        self.db.commit()
        self.update_table()

    def reg_register_button_clicked(self):
        message1 = QMessageBox(MainWindow)
        message1.setIcon(QMessageBox.Warning)
        message1.setWindowTitle("Неправильные данные")

        password1 = self.window.regPassFEdit.text()
        password2 = self.window.regPassSEdit.text()
        login = self.window.regLoginEdit.text()

        if self.check_in_table(login):
            message1.setText("Такой логин уже есть")

        if password1 != password2:
            message1.setText("Пароли не совпадают")

        employee_id = self.window.regIdCB.currentText()
        for record in self.data_table:
            if employee_id == str(record[4]):
                message1.setText("Такой сотрудник уже есть")

        admin_password = self.window.regAdminPassEdit.text()
        role = "Пользователь"

        if self.window.regAdminRadioB.isChecked():
            if len(admin_password) == 0:
                message1.setText("Пароль админа не введён")
            else:
                if admin_password != self.admin_password:
                    message1.setText("Пароль админа не правильный")
                else:
                    role = "Администратор"

        if len(message1.text()) != 0:
            message1.exec_()
            return

        self.insert_into_table(login, password1, role, int(employee_id))

        self.window.tabWidget.setCurrentIndex(0)

    def reg_admin_radio_button_toggled(self):
        if self.window.regAdminRadioB.isChecked():
            self.window.regAdminPassEdit.setVisible(True)
            self.window.label_6.setVisible(True)
        else:
            self.window.regAdminPassEdit.setVisible(False)
            self.window.label_6.setVisible(False)

    def enter_go_to_reg_button_clicked(self):
        self.window.tabWidget.setCurrentIndex(1)

    def reg_return_to_enter_button_clicked(self):
        self.window.tabWidget.setCurrentIndex(0)


class From1(QtWidgets.QMainWindow):
    def __init__(self, login, role, employee_id):
        super(From1, self).__init__()
        self.window = Ui_MainWindow()
        self.window.setupUi(self)

        self.login, self.role, self.employee_id = login, role, employee_id

        headers = {
            "Сотрудники": ["Код сотрудника", "Имя", "Фамилия", "Отчество", "Опыт работы", "Должность", "Отдел",
                           "Учавствовал в инвентаризации"],
            "Должности": ["Название долнжости", "Описание", "Требуемый опыт"],
            "Отделы": ["Номер отдела", "Название отдела"],
            "Имущество": ["Код предмета", "Поставщик", "Ответственной лицо", "Расположение", "Тип", "Состояние",
                          "Стоимость"],
            "Поставщики": ["Наименование организации", "Телефон"],
            "Расположения": ["Корпус", "Этаж", "Кабинет"],
            "Типы имущества": ["Тип", "Описание"],
            "ИКомиссии": ["Номер комиссии", "Дата начала", "Дата окончания", "Причина проверки", "Тип проверки"]
        }
        self.table = Table(self.window, headers)

        self.__deleted_records = [[], ""]

        self.__data = []
        self.form_initialization()

    def form_initialization(self):
        tabs = ["Сотрудники", "Должности", "Отделы", "Имущество", "Поставщики", "Расположения",
                "Типы имущества", "ИКомиссии"]
        for tab in tabs:
            temp = QtWidgets.QWidget()
            temp.setGeometry(QtCore.QRect(0, 0, 0, 0))
            self.window.tablesTabs.addTab(temp, tab)

        staff_table = self.table.get_table("Сотрудники")
        fio = ""
        for record in staff_table:
            if record[0] == self.employee_id:
                fio = record[2] + ' ' + record[1] + ' ' + record[3]

        self.window.FIOLabel.setText(self.window.FIOLabel.text() + fio)
        self.window.RoleLabel.setText(self.window.RoleLabel.text() + self.role)
        self.window.LoginLabel.setText(self.window.LoginLabel.text() + self.login)
        self.window.label_25.setText(self.window.label_25.text() + f'({self.employee_id})')

        self.table.display_table("Сотрудники")
        self.input_tabs_enable_tab(0)
        self.filter_table_enable(0)

        self.window.tableWidget.resizeColumnsToContents()

        self.reset()
        #
        self.window.tableWidget.clicked.connect(self.save_cell)
        self.window.tableWidget.itemChanged.connect(self.update_table)
        #
        self.window.filterCombobox.currentIndexChanged.connect(self.filter_combobox_changed)
        #
        self.window.additionCombobox.currentIndexChanged.connect(self.addition_combobox_changed)
        self.window.additionAddButton.clicked.connect(self.addition_add_button_click)
        self.window.tablesTabs.currentChanged.connect(self.table_tabs_changed)
        #
        self.window.filterStaffDecRadioB_2.toggled.connect(self.filter_staff)
        self.window.filterStaffIncRadioB.toggled.connect(self.filter_staff)
        self.window.filterStaffDecRadioB.toggled.connect(self.filter_staff)
        self.window.filterStaffPostsCB.currentIndexChanged.connect(self.filter_staff)
        self.window.filterStaffDepartCB.currentIndexChanged.connect(self.filter_staff)
        self.window.filterStaffExpFromSpinB.valueChanged.connect(self.filter_staff)
        self.window.filterStaffExpToSpinB.valueChanged.connect(self.filter_staff)
        self.window.filterStaffSortNSearchCB.currentIndexChanged.connect(self.filter_staff)
        self.window.filterStaffSeachEdit.textChanged.connect(self.filter_staff)
        #
        self.window.filterPropDecRadioB.toggled.connect(self.filter_property)
        self.window.filterPropFromSpinBox.textChanged.connect(self.filter_property)
        self.window.filterPropToSpinBox.textChanged.connect(self.filter_property)
        self.window.filterPropIncRadioB.toggled.connect(self.filter_property)
        self.window.filterPropDecRadioB_2.toggled.connect(self.filter_property)
        self.window.filterPropTypeCB.currentIndexChanged.connect(self.filter_property)
        self.window.filterPropStateCB.currentIndexChanged.connect(self.filter_property)
        self.window.filterPropProvidersCB.currentIndexChanged.connect(self.filter_property)
        self.window.filterPropLocationCB.currentIndexChanged.connect(self.filter_property)
        self.window.filterPropSearchEdit.textChanged.connect(self.filter_property)
        self.window.filterPropSearchCB.currentIndexChanged.connect(self.filter_property)

        self.window.filterProvidersSearchEdit.textChanged.connect(self.filter_providers)
        self.window.filterProvidersSortCB.currentIndexChanged.connect(self.filter_providers)

        self.window.filterDeleteButton.clicked.connect(self.filter_delete_button_clicked)

        self.window.removalDeleteButton.clicked.connect(self.removal_delete_button_clicked)

        self.window.removalReturnButton.clicked.connect(self.removal_return_button_clicked)

        self.window.option_tabs.setCurrentIndex(0)
        self.window.option_tabs.currentChanged.connect(self.option_tabs_changed)
        self.window.printInWordButton.clicked.connect(self.print_in_word_button_clicked)
        self.window.printInExcel.clicked.connect(self.print_in_excel_button_clicked)

        self.window.commissionTypeCB.currentTextChanged.connect(self.commission_type_cb_changed)
        self.window.commissionResponsPersonEdit.setVisible(False)
        self.window.commissionLocationCB.setVisible(False)
        self.window.commissionResponsPersonL.setVisible(False)
        self.window.commissionLocationL.setVisible(False)

    def commission_type_cb_changed(self):
        if self.window.commissionReasonCB.count() > 2:
            self.window.commissionReasonCB.removeItem(2)
        current_text = self.window.commissionTypeCB.currentText()
        self.window.commissionResponsPersonEdit.setVisible(False)
        self.window.commissionLocationCB.setVisible(False)
        self.window.commissionResponsPersonL.setVisible(False)
        self.window.commissionLocationL.setVisible(False)
        if current_text == "По расположению":
            self.window.commissionLocationCB.setVisible(True)
            self.window.commissionLocationL.setVisible(True)
        if current_text == "По сотруднику":
            self.window.commissionResponsPersonEdit.setVisible(True)
            self.window.commissionResponsPersonL.setVisible(True)
            self.window.commissionReasonCB.addItem("Увольнение сотрудника")

    def print_in_word_button_clicked(self):
        record = self.detect_commission_record()
        if record is None:
            return

        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        file_name, _ = QFileDialog.getSaveFileName(self, "Выберите расположение файла", f"result{record[0]}.doc",
                                                   "Word (*.doc)", options=options)
        if len(file_name) == 0:
            self.window.message.setText("Вы не выбрали место сохранения!")
            self.window.message.exec_()
            return
        s = file_name.split(".")
        s = s[:-1]
        file_name = ".".join(s)
        file_name = file_name + ".doc"

        doc = DocxTemplate("text.doc")
        context = {
            'organization_name': "ОАО \"Инвент\"",
            'structure_department_name': "ПО-32",
            'start_date': f'{record[1].strftime("%d-%m-%Y")}',
            'end_date': f'{record[2].strftime("%d-%m-%Y")}',
            'now_date': f'{datetime.datetime.now().strftime("%d-%m-%Y")}',
            'reason': f'{record[3]}',
            'type_of': f'{record[4]}',
        }
        doc.render(context)
        doc.save(file_name)

        sql = f'''SELECT имущество2.idИмущество, поставщики.НаименованиеОрганизации, имущество2.Состояние, имущество2.Стоимость
         FROM поставщики INNER JOIN имущество2 ON поставщики.idПоставщики = имущество2.idПоставщики
            WHERE имущество2.idИКомиссии = {record[0]}'''
        print(record[0])
        table = self.table.select_table(sql)
        print(table)
        sql = f'''SELECT сотрудники.Фамилия, сотрудники.Имя, сотрудники.Отчество 
        FROM составкомиссий INNER JOIN сотрудники on составкомиссий.idСотрудники = сотрудники.idСотрудники 
        WHERE idКомиссии = {record[0]}'''
        table2 = self.table.select_table(sql)

        document = Document(file_name)

        if len(table2) != 0:
            document.add_paragraph("Учавствовашие лица")
            headers = ["Фамилия", "Имя", "Отчество"]
            row_count = len(table2) + 1
            column_count = len(table2[0])
            word_table = document.add_table(rows=row_count, cols=column_count)
            word_table.style = 'Table Grid'
            for i in range(0, 3):
                cell = word_table.rows[0].cells[i]
                cell.text = f"{headers[i]}"
                run = cell.paragraphs[0].runs[0]
                run.font.bold = True

            for i in range(1, row_count):
                for j in range(0, column_count):
                    cell = word_table.rows[i].cells[j]
                    cell.text = f"{table2[i - 1][j]}"

        document.add_paragraph("Проверенное имущество")

        row_count = len(table) + 1
        column_count = len(table[0])
        word_table = document.add_table(rows=row_count, cols=column_count)
        word_table.style = 'Table Grid'
        headers = ["Код предмета", "Поставщик", "Состояние", "Стоимость"]
        for i in range(0, 4):
            cell = word_table.rows[0].cells[i]
            cell.text = f"{headers[i]}"
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
        for i in range(1, row_count):
            for j in range(0, column_count):
                cell = word_table.rows[i].cells[j]
                cell.text = f"{table[i - 1][j]}"
        document.save(file_name)

        message1 = QMessageBox(MainWindow)
        message1.setWindowTitle("Успех!")
        message1.setText("Проверьте расположение, там вас ждёт новый файл")
        message1.setIcon(QMessageBox.Information)
        message1.exec_()

    def print_in_excel_button_clicked(self):
        record = self.detect_commission_record()
        if record is None:
            return

        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        file_name, _ = QFileDialog.getSaveFileName(self, "Выберите расположение файла", f"result{record[0]}.xlsx",
                                                   "Excel (*.xlsx)", options=options)
        if len(file_name) == 0:
            self.window.message.setText("Вы не выбрали место сохранения!")
            self.window.message.exec_()
            return
        s = file_name.split(".")
        s = s[:-1]
        file_name = ".".join(s)
        file_name = file_name + ".xlsx"

        sql = f'''SELECT SUM(имущество2.Стоимость), COUNT(имущество2.idИмущество)
         FROM имущество2 
         WHERE Состояние = \'Утерян\' AND имущество2.idИКомиссии = {record[0]}'''
        t = self.table.select_table(sql)
        try:
            s = int(t[0][0])
        except TypeError:
            s = 0
        count = t[0][1]
        data1 = {
            "Номер комиссии": [record[0]],
            "Дата начала проведения": [record[1].strftime("%d-%m-%Y")],
            "Дата окончания проведения": [record[2].strftime("%d-%m-%Y")],
            "Причина проверки": [record[3]],
            "Тип проверки": [record[4]],
            "Количество потерянных предметов": [count],
            "Суммарная стоимость утерянных предметов": [s]
        }
        sql = f'''SELECT имущество2.idИмущество, имущество2.Стоимость, имущество2.ОтветственноеЛицо, имущество2.Состояние
         FROM имущество2 
         WHERE имущество2.Состояние = \'Утерян\' AND имущество2.idИКомиссии = {record[0]}; '''
        property_table = self.table.select_table(sql)
        data2 = {
            "Код предмета": [record[0] for record in property_table],
            "Стоимость": [record[1] for record in property_table],
            "Ответственные лица": [record[2] for record in property_table],
            "Состояние": [record[3] for record in property_table]
        }
        my_wb = openpyxl.Workbook()
        my_wb.active.title = "Обобщённые данные"
        my_wb.create_sheet(index=1, title="Предметы")
        my_wb.active.column_dimensions['A'].width = 20
        my_wb.active.column_dimensions['B'].width = 27
        my_wb.active.column_dimensions['C'].width = 32
        my_wb.active.column_dimensions['D'].width = 32
        my_wb.active.column_dimensions['E'].width = 28
        my_wb.active.column_dimensions['F'].width = 40
        my_wb.active.column_dimensions['G'].width = 50

        header_border = Border(top=Side(border_style='thick', color='FF000000'),
                               right=Side(border_style='thick', color='FF000000'),
                               bottom=Side(border_style='thick', color='FF000000'),
                               left=Side(border_style='thick', color='FF000000'))

        simple_border = Border(top=Side(border_style='thin', color='FF000000'),
                               right=Side(border_style='thin', color='FF000000'),
                               bottom=Side(border_style='thin', color='FF000000'),
                               left=Side(border_style='thin', color='FF000000'))

        for i, key in enumerate(data1.keys()):
            my_wb.active.cell(row=1, column=i + 1).value = key
            my_wb.active.cell(row=1, column=i + 1).font = Font(size=10, name='Comic Sans MS')
            for j, value in enumerate(data1[key]):
                my_wb.active.cell(row=1 + j + 1, column=i + 1).value = value
                my_wb.active.cell(row=1 + j + 1, column=i + 1).font = Font(size=15, name='Comic Sans MS')
                my_wb.active.cell(row=1 + j + 1, column=i + 1).border = simple_border
            my_wb.active.cell(row=1, column=i + 1).border = header_border

        my_wb.active = 1
        my_wb.active.column_dimensions['A'].width = 15
        my_wb.active.column_dimensions['B'].width = 15
        my_wb.active.column_dimensions['C'].width = 40
        my_wb.active.column_dimensions['D'].width = 15
        for i, key in enumerate(data2.keys()):
            my_wb.active.cell(row=1, column=i + 1).value = key
            my_wb.active.cell(row=1, column=i + 1).font = Font(size=10, name='Comic Sans MS')
            for j, value in enumerate(data2[key]):
                my_wb.active.cell(row=1 + j + 1, column=i + 1).value = value
                my_wb.active.cell(row=1 + j + 1, column=i + 1).font = Font(size=12, name='Comic Sans MS')
                my_wb.active.cell(row=1 + j + 1, column=i + 1).border = simple_border
            my_wb.active.cell(row=1, column=i + 1).border = header_border

        my_wb.active = 0

        my_wb.save(file_name)

        message1 = QMessageBox(MainWindow)
        message1.setWindowTitle("Успех!")
        message1.setText("Проверьте расположение, там вас ждёт новый файл")
        message1.setIcon(QMessageBox.Information)
        message1.exec_()

    def detect_commission_record(self):
        if len(self.window.tableWidget.selectedItems()) > 1:
            self.window.message.setText("Выберите только одну запись")
            self.window.message.exec_()
            return None
        table = self.table.get_table("ИКомиссии", with_all_ids=True)
        selected_item = self.window.tableWidget.selectedItems()[0]
        for i, record in enumerate(table):
            if i == selected_item.row():
                return record

    def option_tabs_changed(self):
        if self.window.option_tabs.currentIndex() == 0:
            if self.window.additioninputTabs.currentIndex() in [0, 1, 2, 3, 4]:
                self.window.tablesTabs.setCurrentIndex(self.window.additioninputTabs.currentIndex())
            elif self.window.additioninputTabs.currentIndex() == 5:
                self.window.tablesTabs.setCurrentIndex(6)
        elif self.window.option_tabs.currentIndex() == 2:
            if self.window.filterTabs.currentIndex() == 0:
                self.window.tablesTabs.setCurrentIndex(0)
            elif self.window.filterTabs.currentIndex() == 1:
                self.window.tablesTabs.setCurrentIndex(3)
            elif self.window.filterTabs.currentIndex() == 2:
                self.window.tablesTabs.setCurrentIndex(4)
        if self.window.option_tabs.currentIndex() == 3:
            self.window.tablesTabs.setCurrentIndex(7)

    def save_cell(self):
        global previous_cell
        row = self.window.tableWidget.currentItem().row()
        column = self.window.tableWidget.currentItem().column()
        text = self.window.tableWidget.currentItem().text()
        previous_cell = (row, column, text)

    # I HATE THIS METHOD. I WON'T LOOK HERE EVER. EVEN If THERE IS UNNECESSARY THINGS
    def update_table(self):
        global previous_cell
        current_frame = inspect.currentframe()
        caller_frame = current_frame.f_back
        code_obj = caller_frame.f_code
        code_obj_name = code_obj.co_name
        if code_obj_name != "<module>":
            return

        k = 0
        if self.window.tablesTabs.currentIndex() != 0 and \
                self.window.tablesTabs.currentIndex() != 3 and self.window.tablesTabs.currentIndex() != 2:
            k = 1
        table_name = self.table.get_table_name(self.window.tablesTabs.currentIndex())
        table = self.table.get_table(table_name, with_id=True)
        for i, record in enumerate(table):
            if i == previous_cell[0]:
                record = list(record)
                row_index = previous_cell[1] + k
                new_value = self.window.tableWidget.item(previous_cell[0], previous_cell[1]).text()
                record[row_index] = new_value
                identifier = record[0]

                for j, value in enumerate(record):
                    record[j] = str(value)

                if table_name == "Сотрудники":
                    result = self.validate_staff_addition(*record[1:], fk_is_not_number=True)
                    if result is not None:
                        record = result[0]
                        if row_index == 5:
                            self.window.tableWidget.item(previous_cell[0], previous_cell[1]).setText(result[1])
                        elif row_index == 6:
                            self.window.tableWidget.item(previous_cell[0], previous_cell[1]).setText(result[2])
                        new_value = record[row_index - 1]
                    else:
                        record = result

                elif table_name == "Должности":
                    record = self.validate_posts_addition(*record[1:])
                elif table_name == "Отделы":
                    record = self.validate_department_addition(*record)
                elif table_name == "Имущество":
                    errors = []
                    if row_index == 1:
                        try:
                            if not self.table.is_id_in_table("Поставщики", int(new_value)):
                                errors.append("Такой код сотрудника не существует")
                        except ValueError:
                            errors.append("В это поле вводится код поставщика")

                    elif row_index == 2:
                        try:
                            if not self.table.is_id_in_table("Сотрудники", int(new_value)):
                                errors.append("Такой код сотрудника не существует")
                        except ValueError:
                            errors.append("В это поле вводится код сотрудника")

                    elif row_index == 3:
                        locations_table = self.table.get_table("Расположения", with_id=True)
                        location = 0
                        for rec in locations_table:
                            loc = str(rec[1]) + str(rec[2]) + str(rec[3])
                            if new_value == loc:
                                location = rec[0]
                                new_value = rec[0]
                                break
                        if location == 0:
                            errors.append("Нет такого расположения")
                    elif row_index == 4:
                        types_table = self.table.get_table("Типы имущества", with_id=True)
                        new_value = "".join([letter.lower() for letter in new_value])
                        t = 0
                        for rec in types_table:
                            ty = "".join([letter.lower() for letter in rec[1]])
                            if new_value == ty:
                                t = rec[0]
                                new_value = rec[0]
                                break
                        if t == 0:
                            errors.append("Нет такого типа")
                    elif row_index == 5:
                        error = self.string_validate(new_value, "Состояние")
                        if error != "":
                            errors.append(error)
                    elif row_index == 6:
                        error = self.int_validate(new_value, "Стоимость")
                        if error != "":
                            errors.append(error)
                    if len(errors) != 0:
                        self.window.message.setText(errors[0])
                        self.window.message.exec_()
                        self.window.tableWidget.item(previous_cell[0], previous_cell[1]).setText(previous_cell[2])
                        return
                elif table_name == "Поставщики":
                    record = self.validate_providers_addition(*record[1:])
                elif table_name == "Типы имущества":
                    table_name = "Типы"
                    record = self.validate_types_addition(*record[1:])
                if record is None:
                    self.window.tableWidget.item(previous_cell[0], previous_cell[1]).setText(previous_cell[2])
                    return
                record = list(record)
                record.insert(0, identifier)
                self.table.update_record_in_table(table_name, identifier, row_index, new_value)
                if table_name == "Сотрудники":
                    staff_table = self.table.get_table("Сотрудники")
                    for employee in staff_table:
                        if str(employee[0]) == str(self.employee_id):
                            self.window.FIOLabel.setText(
                                self.window.FIOLabel.text() + employee[2] + ' ' + employee[1] + ' ' + employee[3])
                if table_name == "Типы":
                    table_name = "Типы имущества"
                self.table.display_table(table_name)
                break

    def filter_combobox_changed(self):
        self.filter_table_enable(self.window.filterCombobox.currentIndex())
        index = self.window.filterCombobox.currentIndex()
        if index == 0:
            self.table.display_table("Сотрудники")
            self.window.tablesTabs.setCurrentIndex(0)
        elif index == 1:
            self.table.display_table("Имущество")
            self.window.tablesTabs.setCurrentIndex(3)
        elif index == 2:
            self.table.display_table("Поставщики")
            self.window.tablesTabs.setCurrentIndex(4)

    def filter_table_enable(self, tabindex):
        for index in range(0, self.window.filterTabs.count()):
            self.window.filterTabs.setTabEnabled(index, False)

        self.window.filterTabs.setTabEnabled(tabindex, True)
        self.window.filterTabs.setCurrentIndex(tabindex)

    def removal_return_button_clicked(self):
        table_name = self.__deleted_records[1]
        records = self.__deleted_records[0]

        if len(records) == 0:
            self.window.message.setText("Буфер удалённых записей пуст")
            self.window.message.exec_()
            return

        with_id = False
        if table_name == "Имущество" or table_name == "Отделы":
            with_id = True
        else:
            for i in range(0, len(records)):
                records[i] = records[i][1:]

        for record in records:
            self.table.insert_into_table(table_name, *record, with_id=with_id)

        if table_name == "Типы":
            table_name = "Типы имущества"
        self.table.display_table(table_name)

        self.__deleted_records = [[], ""]

        message1 = QMessageBox(MainWindow)
        message1.setWindowTitle("Успех!")
        message1.setText("Записи возвращены")
        message1.setIcon(QMessageBox.Information)
        message1.exec_()

    def removal_delete_button_clicked(self):
        records = []

        index = self.window.tablesTabs.currentIndex()
        table_name = self.table.get_table_name(index)
        table = self.table.get_table(table_name, with_all_ids=True)
        if table_name == "Типы имущества":
            table_name = "Типы"
        for item in self.window.tableWidget.selectedItems():
            for i, record in enumerate(table):
                if i == item.row():
                    records.append(record)
        self.delete_records(table_name, records)

    def get_answers(self, table_name, questions, error_text, row_index, input_mode, self_value, self_id_error):
        answers = []
        table = self.table.get_table(table_name, with_all_ids=True)
        for i, question in enumerate(questions):
            answered = False
            while answered is False:
                dialog = QtWidgets.QInputDialog()
                if input_mode == 1:
                    answer = dialog.getInt(MainWindow, "Ввод данных", question)
                else:
                    answer = dialog.getText(MainWindow, "Ввод данных", question)
                if answer[1] is False:
                    reply = QMessageBox.question(MainWindow, "Что вы хотите сделать?",
                                                 "Если вы закроете это окно, то все оставшиеся записи удалятся."
                                                 "Хотите удалить все оставшиеся записи? (если нажмёте нет, то удалится "
                                                 "только текущая запись)",
                                                 QMessageBox.Yes | QMessageBox.No | QMessageBox.Cancel)
                    if reply == QMessageBox.Yes:
                        return answers
                    elif reply == QMessageBox.No:
                        break
                answer = answer[0]
                flag = False
                for record in table:
                    value = record[row_index]
                    if input_mode == 0:
                        value = value[0].upper() + value[1:].lower()
                        if len(answer) == 1:
                            answer = answer[0].upper()
                        elif len(answer) > 1:
                            answer = answer[0].upper() + answer[1:].lower()
                    if value == answer:
                        if value == self_value:
                            QMessageBox.information(MainWindow, "Неправильные данные", self_id_error)
                            flag = True
                            break
                        answered = True
                        answers.append((i, record[0]))
                        break
                if flag:
                    continue
                if answered is False:
                    QMessageBox.information(MainWindow, "Неправильные данные", error_text)
        return answers

    def delete_records(self, table_name, records):
        if table_name == "Типы":
            button_reply = QMessageBox.question(MainWindow, "Удаление данных НЕОБРАТИМО",
                                                "Если вы удалите тип(ы) имущества - удалятся все записи об имуществе,"
                                                " которе имеет соответствующий тип(ы). Вы согласны с этим?",
                                                QMessageBox.Yes, QMessageBox.No)
            if button_reply == QMessageBox.No:
                return
        elif table_name == "Расположения":
            QMessageBox.information(MainWindow, "Действие отмененно", "Нельзя удалять расположения")
            return
        elif table_name == "Поставщики":
            QMessageBox.information(MainWindow, "Действие отмененно", "Нельзя удалять поставщиков. Тут указаны те"
                                                                      " предприятия, которые хоть раз поставили"
                                                                      " что-либо.")
            return
        elif table_name == "Отделы":
            values = []

            staff_table = self.table.get_table("Сотрудники", with_all_ids=True)

            departments = []
            for depart in records:
                for employee in staff_table:
                    if depart[0] == employee[6]:
                        departments.append(depart)
                        break

            for depart in departments:
                button_reply = QMessageBox.question(MainWindow, "Удаляете отдел?",
                                                    f"Хотите отправить сотрудников этого отдела ({depart[1]}) в другой?"
                                                    " (при нажати кнопки \"Нет\" все сотрудники выбранного отдела"
                                                    " будут удалены!)", QMessageBox.Yes, QMessageBox.No)
                if button_reply == QMessageBox.No:
                    continue

                this_depart_employees = []
                for employee in staff_table:
                    if employee[6] == depart[0]:
                        this_depart_employees.append(employee)

                questions = []
                for employee in this_depart_employees:
                    string = f"Для сотрудника {employee[2]} {employee[1]} {employee[3]} (код {employee[0]})" \
                             f" требуется указать новый отдел."
                    questions.append(string)

                answers = self.get_answers("Отделы", questions, "Нет такого отдела", 1, 0, depart[1],
                                           "Вы ввели тот же отдел, что удаляете!")

                for answer in answers:
                    values.append(("Сотрудники", this_depart_employees[answer[0]][0], 6, answer[1]))

            button_reply = QMessageBox.question(MainWindow, "Удаление данных НЕОБРАТИМО",
                                                "Вы уверены, что хотите удалить данные?", QMessageBox.Yes,
                                                QMessageBox.No)
            if button_reply == QMessageBox.No:
                return

            for value in values:
                self.table.update_record_in_table(value[0], value[1], value[2], value[3])
        elif table_name == "Сотрудники":
            current_frame = inspect.currentframe()
            caller_frame = current_frame.f_back
            code_obj = caller_frame.f_code
            code_obj_name = code_obj.co_name
            if code_obj_name != "add_record_to_commissions":
                message1 = QMessageBox(MainWindow)
                message1.setWindowTitle("Трубется действие")
                message1.setText("Прежде чем удалить сотрудника(ов), необходимо ввести результаты его инвентаризаци.")
                message1.setIcon(QMessageBox.Information)
                message1.exec_()

            values = []
            property_table = self.table.get_table("Имущество", with_all_ids=True)
            employees = []
            for employee in records:
                for item in property_table:
                    if item[2] == employee[0]:
                        employees.append(employee)
                        break

            for employee in employees:
                print(0)
                if str(employee[0]) == str(self.employee_id):
                    records.remove(employee)
                    message2 = QMessageBox(MainWindow)
                    message2.setWindowTitle("Предупреждение")
                    message2.setText("Сотрудник, чья сессия сейчас запущена удалён не будет.")
                    message2.setIcon(QMessageBox.Information)
                    message2.exec_()
                    continue
                self.__data = []
                ok = False
                flag = False
                print(1)
                if code_obj_name != "add_record_to_commissions":
                    while not ok:
                        start_date, end_date, ok = DateDialog.get_date()
                        if not ok:
                            reply = QMessageBox.question(MainWindow, "Уверены?",
                                                         "Если не введёте инвентаризацию по сотруднику, он не"
                                                         "будет удалён. Вы согласны с этим?", QMessageBox.Yes,
                                                         QMessageBox.No)
                            if reply == QMessageBox.No:
                                continue
                            elif reply == QMessageBox.Yes:
                                flag = True
                                break
                        self.add_record_to_commissions(start_date, end_date, "Увольнение сотрудника", "По сотруднику",
                                                       employee[0], "111")

                    if flag:
                        records.remove(employee)
                        continue

                print(2)

                button_reply = QMessageBox.question(MainWindow, "Удаляете сотрудника?",
                                                    f"Этот сотрудник (код {employee[0]})"
                                                    f" ответственен за некоторое имущество."
                                                    " Хотите назначить новых ответственных лиц?"
                                                    " (при нажати кнопки \"Нет\" всё имущество,"
                                                    " за которое ответственен"
                                                    " этот сотрудник"
                                                    " будет удалено!)", QMessageBox.Yes, QMessageBox.No)
                if button_reply == QMessageBox.No:
                    continue

                print(3)

                this_employee_items = []
                for item in property_table:
                    if employee[0] == item[2]:
                        this_employee_items.append(item)

                questions = []
                for item in this_employee_items:
                    string = f"Тип: {item[4]}; расположение: {item[3]}; код: {item[0]}." \
                             f" Укажите новое ответственное лицо (код сотрудника)."
                    questions.append(string)

                answers = self.get_answers("Сотрудники", questions, "Нет такого сотрудника", 0, 1, employee[0],
                                           "Вы ввели того же сотрудника, которого удаляете!")

                for answer in answers:
                    values.append(("Имущество", this_employee_items[answer[0]][0], 2, answer[1]))

            if len(records) == 0:
                return

            button_reply = QMessageBox.question(MainWindow, "Удаление данных НЕОБРАТИМО",
                                                "Вы уверены, что хотите удалить данные?", QMessageBox.Yes,
                                                QMessageBox.No)
            if button_reply == QMessageBox.No:
                return

            for value in values:
                self.table.update_record_in_table(value[0], value[1], value[2], value[3])
        elif table_name == "Должности":
            values = []
            staff_table = self.table.get_table("Сотрудники", with_all_ids=True)

            posts = []
            for post in records:
                for employee in staff_table:
                    if post[0] == employee[5]:
                        posts.append(post)
                        break

            for post in posts:
                button_reply = QMessageBox.question(MainWindow, "Удаляете должность?",
                                                    f"Хотите назначить сотрудников, имеющих эту должность ({post[1]})"
                                                    f" на другую?"
                                                    " (при нажати кнопки \"Нет\" все сотрудники выбранной должности"
                                                    " будут удалены!)", QMessageBox.Yes, QMessageBox.No)
                if button_reply == QMessageBox.No:
                    continue

                this_post_employees = []
                for employee in staff_table:
                    if employee[5] == post[0]:
                        this_post_employees.append(employee)

                questions = []
                for employee in this_post_employees:
                    string = f"Для сотрудника {employee[2]} {employee[1]} {employee[3]} (код {employee[0]})" \
                             f" требуется указать новую должность."
                    questions.append(string)

                answers = self.get_answers("Должности", questions, "Нет такой должности", 1, 0, post[1],
                                           "Вы ввели ту же должность, что удаляете!")

                for answer in answers:
                    values.append(("Сотрудники", this_post_employees[answer[0]][0], 5, answer[1]))
            button_reply = QMessageBox.question(MainWindow, "Удаление данных НЕОБРАТИМО",
                                                "Вы уверены, что хотите удалить данные?", QMessageBox.Yes,
                                                QMessageBox.No)
            if button_reply == QMessageBox.No:
                return

            for value in values:
                self.table.update_record_in_table(value[0], value[1], value[2], value[3])
        elif table_name == "ИКомиссии":
            QMessageBox.information(MainWindow, "Действие отмененно", "Нельзя удалять инвентаризационные комиссии")
            return

        if table_name not in ("Сотрудники", "Отделы", "Должности"):
            button_reply = QMessageBox.question(MainWindow, "Удаление данных НЕОБРАТИМО",
                                                "Вы уверены, что хотите удалить данные?", QMessageBox.Yes,
                                                QMessageBox.No)
            if button_reply == QMessageBox.No:
                return

        self.__deleted_records[0] = []
        self.__deleted_records[1] = ""

        for record in records:
            self.table.delete_record_from_table(record[0], table_name)
            self.__deleted_records[0].append(record)
            self.__deleted_records[1] = table_name

        if table_name == "Типы":
            table_name = "Типы имущества"

        self.table.update_table(table_name)
        if table_name == "Сотрудники":
            self.table.update_table("Имущество")
        self.table.display_table(table_name)

        self.reset()
        message1 = QMessageBox(MainWindow)
        message1.setWindowTitle("Успех!")
        message1.setText("Записи удалены")
        message1.setIcon(QMessageBox.Information)
        message1.exec_()

    def filter_delete_button_clicked(self):
        filters = {
            "Сотрудники": self.filter_staff,
            "Имущество": self.filter_property,
            "Поставщики": self.filter_providers
        }
        records = []

        if self.window.filterTabs.currentIndex() == 2:
            table_name = "Поставщики"
        elif self.window.filterTabs.currentIndex() == 1:
            table_name = "Имущество"
        else:
            table_name = "Сотрудники"
        table = filters[table_name]()
        for record in table:
            for rec in self.table.get_table(table_name, with_all_ids=True):
                if rec[0] == record[0]:
                    record = rec
                    break
            records.append(record)
        self.delete_records(table_name, records)

    def filter_providers(self):
        table = self.table.get_table("Поставщики")
        table_with_ids = self.table.get_table("Поставщики", with_id=True)

        self.table.display_table("Поставщики", table)
        #
        if len(self.window.filterProvidersSearchEdit.text()) != 0:
            table = self.table.left_found_records(table, self.window.filterProvidersSearchEdit.text(),
                                                  self.window.filterProvidersSortCB.currentIndex() + 1)
            table_with_ids = self.table.left_found_records(table_with_ids, self.window.filterProvidersSearchEdit.text(),
                                                           self.window.filterProvidersSortCB.currentIndex() + 2)
            self.table.color_by_matching_with_another_table(table, (90, 200, 90))

        return table_with_ids

    def filter_property(self):
        table = self.table.get_table("Имущество")

        t = self.window.filterPropTypeCB.currentText()
        location = self.window.filterPropLocationCB.currentText()
        state = self.window.filterPropStateCB.currentText()
        provider = self.window.filterPropProvidersCB.currentText()

        if t != "Все":
            table = self.table.left_only_one_value(table, t, 5)

        if location != "Все":
            table = self.table.left_only_one_value(table, location, 4)

        if state != "Все":
            table = self.table.left_only_one_value(table, state, 6)

        if provider != "Все":
            table = self.table.left_only_one_value(table, provider, 2)

        bottom_value = self.window.filterPropFromSpinBox.value()
        top_value = self.window.filterPropToSpinBox.value()

        table = self.table.left_only_from_to_records(bottom_value, top_value, table, 7)

        if self.window.filterPropDecRadioB.isChecked():
            table = self.table.descending_sort(table, 7)
        if self.window.filterPropIncRadioB.isChecked():
            table = self.table.ascending_sort(table, 7)

        self.table.display_table("Имущество", table)

        if len(self.window.filterPropSearchEdit.text()) != 0:
            table = self.table.convert_table_to_string(table)
            table = self.table.left_found_records(table, self.window.filterPropSearchEdit.text(),
                                                  self.window.filterPropSearchCB.currentIndex() + 2)
            self.table.color_by_matching_with_another_table(table, (90, 200, 90))

        return table

    def filter_staff(self):
        current_frame = inspect.currentframe()
        caller_frame = current_frame.f_back
        code_obj = caller_frame.f_code
        code_obj_name = code_obj.co_name
        if code_obj_name == "reset":
            return

        table = self.table.get_table("Сотрудники")

        post = self.window.filterStaffPostsCB.currentText()
        department = self.window.filterStaffDepartCB.currentText()

        if post != "Все":
            table = self.table.left_only_one_value(table, post, 6)

        if department != "Все":
            table = self.table.left_only_one_value(table, department, 7)

        bottom_value = self.window.filterStaffExpFromSpinB.value()
        top_value = self.window.filterStaffExpToSpinB.value()

        table = self.table.left_only_from_to_records(bottom_value, top_value, table, 5)

        if self.window.filterStaffDecRadioB.isChecked():
            table = self.table.descending_sort(table, self.window.filterStaffSortNSearchCB.currentIndex() + 2)
        if self.window.filterStaffIncRadioB.isChecked():
            table = self.table.ascending_sort(table, self.window.filterStaffSortNSearchCB.currentIndex() + 2)

        self.table.display_table("Сотрудники", table)

        if len(self.window.filterStaffSeachEdit.text()) != 0:
            table = self.table.convert_table_to_string(table)
            table = self.table.left_found_records(table, self.window.filterStaffSeachEdit.text(),
                                                  self.window.filterStaffSortNSearchCB.currentIndex() + 2)
            self.table.color_by_matching_with_another_table(table, (90, 200, 90))

        return table

    def configure_combo_boxes(self):
        def staff_configure():
            posts = self.table.get_table("Должности")
            departments = self.table.get_table("Отделы")

            self.window.filterStaffDepartCB.addItem("Все")
            for depart in departments:
                self.window.staffDepartmentCB.addItem(depart[1])
                self.window.filterStaffDepartCB.addItem(depart[1])

            self.window.filterStaffPostsCB.addItem("Все")

            for post in posts:
                self.window.staffPostCB.addItem(post[0])
                self.window.filterStaffPostsCB.addItem(post[0])

        def property_configure():
            types = set()
            providers = set()
            locations = set()
            states = ["Исправен", "Требуется починка", "Утерян"]
            types_table = self.table.get_table("Типы имущества")
            providers_table = self.table.get_table("Поставщики")
            locations_table = self.table.get_table("Расположения")

            for record in types_table:
                t = record[0]
                types.add(t)
            self.window.filterPropTypeCB.addItem("Все")
            for t in types:
                self.window.propertyTypeCB.addItem(t)
                self.window.filterPropTypeCB.addItem(t)

            for record in providers_table:
                provider = record[0]
                providers.add(provider)
            self.window.filterPropProvidersCB.addItem("Все")
            for provider in providers:
                self.window.propertyProviderCB.addItem(provider)
                self.window.filterPropProvidersCB.addItem(provider)

            self.window.filterPropStateCB.addItem("Все")
            for state in states:
                self.window.propertyStateCB.addItem(state)
                self.window.filterPropStateCB.addItem(state)

            for record in locations_table:
                location = str(record[0]) + str(record[1]) + str(record[2])
                locations.add(location)
            locations = list(locations)
            locations.sort()
            self.window.filterPropLocationCB.addItem("Все")
            for location in locations:
                self.window.propertyLocationCB.addItem(location)
                self.window.filterPropLocationCB.addItem(location)
                self.window.commissionLocationCB.addItem(location)

        def filter_staff_configure():
            fields = ["Имя", "Фамилия", "Отчество", "Опыт работы"]
            for field in fields:
                self.window.filterStaffSortNSearchCB.addItem(field)

        def filter_property_configure():
            fields = ["Поставщики", "Ответственные лица", "Расположения", "Типы"]
            for field in fields:
                self.window.filterPropSearchCB.addItem(field)

        def filter_providers_configure():
            fields = ["Название", "Телефон"]
            for field in fields:
                self.window.filterProvidersSortCB.addItem(field)

        filter_providers_configure()
        filter_property_configure()
        filter_staff_configure()

        staff_configure()

        property_configure()

    def input_tabs_enable_tab(self, tabindex):
        for index in range(0, self.window.additioninputTabs.count()):
            self.window.additioninputTabs.setTabEnabled(index, False)

        self.window.additioninputTabs.setTabEnabled(tabindex, True)
        self.window.additioninputTabs.setCurrentIndex(tabindex)

    def addition_add_button_click(self):
        is_added = False
        if self.window.additionCombobox.currentIndex() == 0:
            is_added = self.add_record_to_staff(self.window.staffFNameEdit.text(), self.window.staffSNameEdit.text(),
                                                self.window.staffPatEdit.text(), self.window.staffExpEdit.text(),
                                                self.window.staffPostCB.currentText(),
                                                self.window.staffDepartmentCB.currentText())
        elif self.window.additionCombobox.currentIndex() == 1:
            is_added = self.add_record_to_posts(self.window.postsTitleEdit.text(),
                                                self.window.typesDiscEdit.toPlainText(),
                                                self.window.postsExpEdit.text())
        elif self.window.additionCombobox.currentIndex() == 2:
            is_added = self.add_record_to_departments(self.window.departmentsIdEdit.text(),
                                                      self.window.departmentsTitleEdit.text())
        elif self.window.additionCombobox.currentIndex() == 3:
            is_added = self.add_record_to_property(self.window.propertyIdEdit.text(),
                                                   self.window.propertyProviderCB.currentText(),
                                                   self.window.propertyEmplIdEdit.text(),
                                                   self.window.propertyLocationCB.currentText(),
                                                   self.window.propertyTypeCB.currentText(),
                                                   self.window.propertyStateCB.currentText(),
                                                   self.window.propertyCostEdit.text())
        elif self.window.additionCombobox.currentIndex() == 4:
            is_added = self.add_record_to_providers(self.window.providersTitleEdit.text(),
                                                    self.window.providersPhoneEdit.text())
        elif self.window.additionCombobox.currentIndex() == 5:
            is_added = self.add_record_to_types(self.window.typesTypeEdit.text(),
                                                self.window.typesDiscEdit.toPlainText())
        elif self.window.additionCombobox.currentIndex() == 6:
            is_added = self.add_record_to_commissions(self.window.commissionstartDate.text(),
                                                      self.window.commissionendDate.text(),
                                                      self.window.commissionReasonCB.currentText(),
                                                      self.window.commissionTypeCB.currentText(),
                                                      self.window.commissionResponsPersonEdit.text(),
                                                      self.window.commissionLocationCB.currentText())
        #############################################################
        if is_added:
            self.table.update_table(self.window.additionCombobox.currentText())
            self.table.display_table(self.window.additionCombobox.currentText())
            self.reset()
            message1 = QMessageBox(MainWindow)
            message1.setWindowTitle("Успех!")
            message1.setText("Запись добавлена")
            message1.setIcon(QMessageBox.Information)
            message1.exec_()

    def addition_combobox_changed(self):
        index = self.window.additionCombobox.currentIndex()
        self.input_tabs_enable_tab(index)
        self.window.tablesTabs.setCurrentIndex(index)
        if index == 5:
            self.table.display_table("Типы имущества")
            self.window.tablesTabs.setCurrentIndex(6)
        elif index == 6:
            self.table.display_table("ИКомиссии")
            self.window.tablesTabs.setCurrentIndex(7)
        else:
            self.table.display_table(self.window.additionCombobox.currentText())
            self.window.tablesTabs.setCurrentIndex(index)

    def table_tabs_changed(self):
        index = self.window.tablesTabs.currentIndex()
        table_name = self.table.get_table_name(index)
        self.table.display_table(table_name)

    @staticmethod
    def string_validate(string, field):
        if len(string) == 0:
            return f"Поле \"{field}\" должно состоять как минимум из 1-го символа"
        elif True in [letter in digits for letter in string]:
            return f"В поле \"{field}\" не может быть чисел"
        elif len(string) > 45:
            return f"Поле \"{field}\" должно быть короче 45-ти символов"
        else:
            return ""

    @staticmethod
    def int_validate(string, field):
        try:
            int(string)
            return ""
        except ValueError:
            return f"Поле \"{field}\" должно быть целочисленным"

    def reset(self):
        self.configure_combo_boxes()
        self.window.staffPatEdit.setText("")
        self.window.staffFNameEdit.setText("")
        self.window.staffSNameEdit.setText("")
        self.window.staffExpEdit.setText("")
        self.window.postsTitleEdit.setText("")
        self.window.postsDicsEdit.setText("")
        self.window.postsExpEdit.setText("")
        self.window.departmentsTitleEdit.setText("")
        self.window.departmentsIdEdit.setText("")
        self.window.propertyCostEdit.setText("")
        self.window.propertyIdEdit.setText("")
        self.window.propertyEmplIdEdit.setText("")
        self.window.providersPhoneEdit.setText("")
        self.window.providersTitleEdit.setText("")
        self.window.typesTypeEdit.setText("")
        self.window.typesDiscEdit.setText("")
        #
        #
        self.window.filterStaffDepartCB.setCurrentIndex(0)
        self.window.filterStaffPostsCB.setCurrentIndex(0)
        self.window.filterStaffSortNSearchCB.setCurrentIndex(0)
        table = self.table.get_table("Сотрудники")
        maximum = max([record[4] for record in table])
        self.window.filterPropFromSpinBox.setValue(0)
        self.window.filterStaffExpToSpinB.setValue(maximum)
        self.window.filterStaffDecRadioB_2.setChecked(True)
        self.window.filterStaffSeachEdit.setText("")
        #
        self.window.filterPropLocationCB.setCurrentIndex(0)
        self.window.filterPropProvidersCB.setCurrentIndex(0)
        self.window.filterPropTypeCB.setCurrentIndex(0)
        self.window.filterPropStateCB.setCurrentIndex(0)
        self.window.filterPropDecRadioB_2.setChecked(True)
        self.window.filterPropFromSpinBox.setValue(0)
        table = self.table.get_table("Имущество")
        maximum = max([record[6] for record in table])
        self.window.filterPropToSpinBox.setValue(maximum)
        self.window.filterPropSearchEdit.setText("")
        #
        self.window.filterProvidersSearchEdit.setText("")
        #
        now = datetime.datetime.now()
        self.window.commissionstartDate.setDate(now)
        self.window.commissionendDate.setDate(now)
        self.window.commissionTypeCB.setCurrentIndex(0)
        self.window.commissionReasonCB.setCurrentIndex(0)

    # Имущество
    def validate_property_addition(self, idd_str, provider_str, responsible_person_str, location_str, t_str, state_str,
                                   cost_str):
        error = ""
        errors = [self.int_validate(idd_str, "Код предмета"), self.int_validate(cost_str, "Стоимость"),
                  self.int_validate(responsible_person_str, "Ответственное лицо")]
        try:
            if self.table.is_id_in_table("Имущество", int(idd_str)):
                errors.append("Такой код предмета уже существует")

            if not self.table.is_id_in_table("Сотрудники", int(responsible_person_str)):
                errors.append("Такой код сотрудника не существует")
        except ValueError:
            pass

        for err in errors:
            if err != "":
                error = err
                break

        self.window.message.setText(error)

        if self.window.message.text() != "":
            self.window.message.exec_()
            return

        t = 0
        types_table = self.table.get_table("Типы имущества", with_id=True)
        provider = 0
        providers_table = self.table.get_table("Поставщики", with_id=True)
        location = 0
        locations_table = self.table.get_table("Расположения", with_id=True)

        idd = int(idd_str)

        cost = int(cost_str)

        state = state_str

        for record in types_table:
            if record[1] == t_str:
                t = record[0]

        for record in locations_table:
            if str(record[1]) + str(record[2]) + str(record[3]) == location_str:
                location = record[0]

        for record in providers_table:
            if record[1] == provider_str:
                provider = record[0]

        responsible_person = int(responsible_person_str)

        result = (idd, provider, responsible_person, location, t, state, cost)

        return result

    def add_record_to_property(self, idd_str, provider_str, responsible_person_str, location_str, t_str, state_str,
                               cost_str):
        record = self.validate_property_addition(idd_str, provider_str, responsible_person_str, location_str, t_str,
                                                 state_str, cost_str)

        if record is None:
            return False

        idd, provider, responsible_person, location, t, state, cost = record

        self.table.insert_into_table("Имущество", idd, provider, responsible_person,
                                     location, t, state, cost, with_id=True)

        return True

    # Поставщики
    def validate_providers_addition(self, org_name_str, phone_str):
        error = ""
        errors = [self.string_validate(org_name_str, "Наименование организации"),
                  self.int_validate(phone_str, "Номер телефона")]

        if len(phone_str) != 12:
            errors.append("Не верный формат номера")

        for err in errors:
            if err != "":
                error = err
                break

        self.window.message.setText(error)

        if self.window.message.text() != "":
            self.window.message.exec_()
            return

        org_name = org_name_str

        phone = phone_str

        result = (org_name, phone)

        return result

    def add_record_to_providers(self, org_name_str, phone_str):
        record = self.validate_providers_addition(org_name_str, phone_str)

        if record is None:
            return False

        org_name, phone = record

        self.table.insert_into_table("Поставщики", org_name, phone)

        return True

    # Типы имущество
    def validate_types_addition(self, property_type_str, disc_str):
        error = ""
        errors = [self.string_validate(property_type_str, "Тип")]

        for err in errors:
            if err != "":
                error = err
                break

        self.window.message.setText(error)

        if self.window.message.text() != "":
            self.window.message.exec_()
            return None

        property_type = property_type_str

        disc = disc_str

        result = (property_type, disc)

        return result

    def add_record_to_types(self, property_type_str, disc_str):
        record = self.validate_types_addition(property_type_str, disc_str)

        if record is None:
            return False

        property_type, disc = record

        self.table.insert_into_table("Типы", property_type, disc)

        return True

    # Сотрудники
    def validate_staff_addition(self, name_str, surname_str, patronymic_str, experience_str, post_str, department_str,
                                fk_is_not_number=False):
        result = ["", ""]
        error = ""
        post = 0
        posts_table = self.table.get_table("Должности", with_id=True)
        department = 0
        depart_table = self.table.get_table("Отделы")

        errors = [self.string_validate(name_str, "Имя"), self.int_validate(experience_str, "Опыт работы"),
                  self.string_validate(surname_str, "Фамилия"), self.string_validate(patronymic_str, "Отчество")]

        if fk_is_not_number:
            post_str = "".join([letter.lower() for letter in post_str])

            for record in posts_table:
                p = "".join([letter.lower() for letter in record[1]])
                if post_str == p:
                    post_str = post_str[0].upper() + post_str[1:]
                    result[0] = post_str
                    post = int(record[0])
                    try:
                        if int(record[3]) > int(experience_str):
                            errors.append("Сотрудник не может занять эту должность, так как его опыт меньше"
                                          f"требуемого на этой должности ({record[3]})")
                    finally:
                        pass
            if len(result[0]) == 0:
                errors.append("Нет такой должности")

            department_str = "".join([letter.lower() for letter in department_str])

            for record in depart_table:
                d = "".join([letter.lower() for letter in record[1]])
                if department_str == d:
                    department_str = department_str[0].upper() + department_str[1:]
                    result[1] = department_str
                    department = int(record[0])
            if len(result[1]) == 0:
                errors.append("Нет такого отдела")
        if not fk_is_not_number:
            for record in posts_table:
                if record[1] == post_str:
                    try:
                        if int(record[3]) > int(experience_str):
                            errors.append("Сотрудник не может занять эту должность, так как его опыт меньше"
                                          f"требуемого на этой должности ({record[3]})")
                    finally:
                        pass
        for err in errors:
            if err != "":
                error = err
                break

        self.window.message.setText(error)

        if self.window.message.text() != "":
            self.window.message.exec_()
            return None

        name = name_str

        surname = surname_str

        patronymic = patronymic_str

        experience = int(experience_str)

        if not fk_is_not_number:
            for record in posts_table:
                if record[1] == post_str:
                    post = int(record[0])

            for record in depart_table:
                if record[1] == department_str:
                    department = int(record[0])

        result.insert(0, (name, surname, patronymic, experience, post, department))

        return result

    def add_record_to_staff(self, name_str, surname_str, patronymic_str, experience_str, post_str, department_str):
        record = self.validate_staff_addition(name_str, surname_str, patronymic_str, experience_str, post_str,
                                              department_str)

        if record is None:
            return False

        name, surname, patronymic, experience, post, department = record[0]

        self.table.insert_into_table("Сотрудники", name, surname, patronymic, experience, post, department)

        return True

    # Должности
    def validate_posts_addition(self, title_str, disc_str, experience_str):
        error = ""
        errors = [self.string_validate(title_str, "Название"), self.int_validate(experience_str, "Требуемый опыт")]

        for err in errors:
            if err != "":
                error = err
                break

        self.window.message.setText(error)

        if self.window.message.text() != "":
            self.window.message.exec_()
            return None

        title = title_str

        experience = int(experience_str)

        disc = disc_str

        result = (title, experience, disc)

        return result

    def add_record_to_posts(self, title_str, disc_str, experience_str):
        record = self.validate_posts_addition(title_str, disc_str, experience_str)

        if record is None:
            return False

        title, experience, disc = record

        self.table.insert_into_table("Должности", title, disc, experience)

        return True

    # Отделы
    def validate_department_addition(self, id_str, title_str):
        error = ""
        errors = [self.string_validate(title_str, "Название отдела"),
                  self.int_validate(self.window.departmentsIdEdit.text(), "Номер отдела")]

        if self.table.is_id_in_table("Отделы", int(id_str)):
            errors.append("Такой номер отдела уже существует")

        for err in errors:
            if err != "":
                error = err
                break

        self.window.message.setText(error)

        if self.window.message.text() != "":
            self.window.message.exec_()
            return None

        title = title_str

        i = int(id_str)

        result = (i, title)

        return result

    def add_record_to_departments(self, id_str, title_str):
        record = self.validate_department_addition(id_str, title_str)

        if record is None:
            return False

        i, title = record

        self.table.insert_into_table("Отделы", i, title, with_id=True)

        return True

    # ИКомиссии
    def add_record_to_commissions(self, start_date_str, end_date_str, reason_str, invent_type_str, responsible_id_str,
                                  location_str):
        def full_invent():
            sql_ = '''SELECT idИмущество, idПоставщики, 
                                CONCAT( '(код ', сотрудники.idСотрудники,') ', сотрудники.Фамилия,' ', сотрудники.Имя,' ', сотрудники.Имя),
                                idРасположения, idТипы, Состояние, Стоимость  
                                FROM имущество INNER JOIN сотрудники on имущество.idОтветственноеЛицо = сотрудники.idСотрудники'''
            return sql_, True, ""

        def location_invent():
            sql_ = f'''SELECT idИмущество, idПоставщики, 
                                CONCAT( '(код ', сотрудники.idСотрудники,') ', сотрудники.Фамилия,' ', сотрудники.Имя,' ', сотрудники.Имя),
                                имущество.idРасположения, idТипы, Состояние, Стоимость  
                                FROM расположения INNER JOIN 
                                (имущество INNER JOIN сотрудники on имущество.idОтветственноеЛицо = сотрудники.idСотрудники)
                                 ON расположения.idРасположения = имущество.idИмущество 
                                 WHERE CONCAT(расположения.Корпус, расположения.Этаж, расположения.Кабинет) = {location_str}'''

            specious_string_ = f" (расположение {location_str})"

            return sql_, True, specious_string_

        def employee_invent():
            errors = [self.int_validate(responsible_id_str, "Ответственное лицо")]
            error = ""
            try:
                if not self.table.is_id_in_table("Сотрудники", int(responsible_id_str)):
                    errors.append("Такой код сотрудника не существует")
            except ValueError:
                pass

            for err in errors:
                if err != "":
                    error = err
                    break

            self.window.message.setText(error)

            if self.window.message.text() != "":
                self.window.message.exec_()
                return "", False, ""

            sql_ = f'''SELECT idИмущество, idПоставщики, 
                                            CONCAT( '(код ', сотрудники.idСотрудники,') ', сотрудники.Фамилия,' ', сотрудники.Имя,' ', сотрудники.Имя),
                                            idРасположения, idТипы, Состояние, Стоимость  
                                            FROM имущество INNER JOIN сотрудники on имущество.idОтветственноеЛицо = сотрудники.idСотрудники
                                            WHERE имущество.idОтветственноеЛицо = {int(responsible_id_str)}'''

            specious_string_ = f" (код сотрудника {responsible_id_str})"

            return sql_, True, specious_string_

        if self.role != "Администратор":
            self.window.message.setText("Добавлять инвентаризационные комиссии может только администратор")
            self.window.message.exec_()
            return False

        start_date_str = start_date_str.split('.')
        start_date_str.reverse()
        date1 = datetime.datetime(int(start_date_str[0]), int(start_date_str[1]), int(start_date_str[2]))
        start_date_str = '-'.join(start_date_str)
        start_date = start_date_str

        end_date_str = end_date_str.split('.')
        end_date_str.reverse()
        date2 = datetime.datetime(int(end_date_str[0]), int(end_date_str[1]), int(end_date_str[2]))
        end_date_str = '-'.join(end_date_str)
        end_date = end_date_str

        ten_days = datetime.timedelta(days=10)

        self.window.message.setText("")
        if date2 < datetime.datetime.now() - ten_days:
            self.window.message.setText("Инвентаризация не могла закончится более 10 дней назад!")
        if date1 > date2:
            self.window.message.setText("Дата окончания не может быть меньше даты начала!")
        if datetime.datetime.now() < date2:
            self.window.message.setText("Инвентаризация не могла закончится позже, чем сегодня!")

        if len(self.window.message.text()) != 0:
            self.window.message.exec_()
            return False

        reason = reason_str

        invent_type = invent_type_str

        if invent_type == "Полная":
            sql, ok, specious_string = full_invent()
        elif invent_type == "По расположению":
            sql, ok, specious_string = location_invent()
        else:
            sql, ok, specious_string = employee_invent()

        if not ok:
            return

        dialog = QtWidgets.QInputDialog()
        employees = []
        ids = []
        staff_table = self.table.get_table("Сотрудники")
        for employee in staff_table:
            ids.append(str(employee[0]))
        emp_count = self.window.commissionSpinBox.value()
        while len(employees) != emp_count:
            employee_id, ok = dialog.getItem(self, "Ввод данных",
                                             f"Выберите код сотрудника",
                                             ids, 0, False)
            if not ok:
                button_reply = QMessageBox.question(MainWindow, "Что вы хотите сделать?",
                                                    "Отменить добавление комиссии?", QMessageBox.Yes,
                                                    QMessageBox.No)
                if button_reply == QMessageBox.Yes:
                    return

            if str(employee_id) == str(responsible_id_str):
                self.window.message.setText("Сотрудник не может учавствовать в инвентаризации себя")
                self.window.message.exec_()
                continue

            if employee_id in employees:
                self.window.message.setText("Такой сотрудник уже выбран")
                self.window.message.exec_()
                continue

            employees.append(employee_id)

        message1 = QMessageBox()
        message1.setText("Сейчас внесите результаты инвентаризации")
        message1.setIcon(QMessageBox.Information)
        message1.exec_()
        message1.setIcon(QMessageBox.Warning)
        property_table = self.table.select_table(sql)
        answers = []
        for record in property_table:
            answered = False
            while answered is False:
                value, answered = dialog.getItem(self, "Ввод данных",
                                                 f"Выберите состояние предмета (код {record[0]}){specious_string}",
                                                 ["Исправен", "Требуется починка", "Утерян"], 0, False)
                if answered is False:
                    button_reply = QMessageBox.question(MainWindow, "Что вы хотите сделать?",
                                                        "Отменить добавление комиссии?", QMessageBox.Yes,
                                                        QMessageBox.No)
                    if button_reply == QMessageBox.Yes:
                        return
                    message1.setText("Тогда вводите данные!")
                    message1.exec_()
                answers.append((record, value))

        self.table.insert_into_table("ИКомиссии", start_date, end_date, reason, invent_type)
        commissions = self.table.get_table("ИКомиссии")
        last = commissions[-1][0]

        for answer in answers:
            self.table.insert_into_table("Имущество2", *answer[0][0:5], answer[1], answer[0][6], last)
            self.table.update_record_in_table("Имущество", answer[0][0], 5, answer[1])

        for employee in employees:
            sql = f'''INSERT INTO `trpo_db`.`составкомиссий` (`idСотрудники`, `idКомиссии`) VALUES ('{int(employee)}', '{int(last)}');'''
            self.table.select_table(sql)
            self.table.db.commit()

        current_frame = inspect.currentframe()
        caller_frame = current_frame.f_back
        code_obj = caller_frame.f_code
        code_obj_name = code_obj.co_name
        if code_obj_name != "delete_records":
            if reason_str == "Увольнение сотрудника":
                print(5, reason)

                button_reply = QMessageBox.question(MainWindow, "Что дальше?",
                                                    "Хотите сразу удалить сотрудника?", QMessageBox.Yes,
                                                    QMessageBox.No)
                if button_reply == QMessageBox.Yes:
                    table = self.table.get_table("Сотрудники", with_all_ids=True)
                    record = []
                    for record_ in table:
                        if record_[0] == int(responsible_id_str):
                            record = record_
                    print("en", record)
                    self.delete_records("Сотрудники", [record])
        return True


class Table:
    def __init__(self, window, headers):
        self.table = window.tableWidget
        self.db = None
        self.tables = {}
        self.tables_with_id = {}
        self.tables_with_all_ids = {}
        self.headers = headers
        self.connect_base()

    def connect_base(self):
        self.db = Mdb.connect('localhost', 'root', '321Ilyxazc', 'trpo_db')
        self.db.set_character_set("utf8")
        self.select_tables()

    def select_table(self, sql):
        cursor = self.db.cursor()
        cursor.execute(sql)
        return cursor.fetchall()

    def select_tables(self):
        staff_sql = '''SELECT сотрудники.idСотрудники, Имя, Фамилия, Отчество, ОпытРаботы, должности.Название, отделы.Название,
                       IF (сотрудники.idСотрудники IN (SELECT составкомиссий.idСотрудники FROM составкомиссий),
                        'Да', 'Нет') AS 'Учавствовал в инвентаризации'
                        FROM отделы INNER JOIN (сотрудники INNER JOIN должности ON
                        сотрудники.idДолжности = должности.idДолжности) 
                        ON отделы.idОтделы = сотрудники.idОтделы
                        ORDER BY сотрудники.idСотрудники;'''
        departments_sql = '''SELECT * FROM отделы'''
        posts_sql = '''SELECT Название, Описание, ТребуемыйОпыт FROM должности'''
        providers_sql = '''SELECT НаименованиеОрганизации, телефон FROM поставщики'''
        types_sql = '''SELECT Наименование, Описание FROM типы'''
        location_sql = '''SELECT Корпус, Этаж, Кабинет FROM расположения'''
        property_sql = '''SELECT имущество.idИмущество, поставщики.НаименованиеОрганизации,
                 concat(сотрудники.Фамилия,' ', сотрудники.Имя,' ', сотрудники.Отчество),
                             concat(расположения.Корпус, расположения.Этаж, расположения.Кабинет)
                              AS Расположение, типы.Наименование,
                              имущество.Состояние, имущество.Стоимость 
                             FROM сотрудники INNER JOIN (типы INNER JOIN (расположения INNER JOIN 
                            (имущество INNER JOIN поставщики on имущество.idПоставщики = поставщики.idПоставщики)
                             ON расположения.idРасположения = имущество.idРасположения) ON типы.idТипы = имущество.idТипы)
                              ON сотрудники.idСотрудники = имущество.idОтветственноеЛицо ORDER BY имущество.idИмущество'''
        commission_sql = '''SELECT икомиссии.idИкомиссии, ДатаНачала, ДатаОкончания, ПричинаПроверки, ТипПроверки FROM икомиссии'''
        property2 = '''SELECT * FROM имущество2'''
        queries = {
            "Сотрудники": staff_sql,
            "Отделы": departments_sql,
            "Должности": posts_sql,
            "Имущество": property_sql,
            "Поставщики": providers_sql,
            "Расположения": location_sql,
            "Типы имущества": types_sql,
            "ИКомиссии": commission_sql,
            "Имущество2": property2
        }
        cursor = self.db.cursor()

        for key in queries.keys():
            cursor.execute(queries[key])
            table = cursor.fetchall()
            self.tables[key] = table

        staff_sql = '''SELECT idСотрудники, Имя, Фамилия, Отчество, ОпытРаботы, должности.Название, отделы.Название
                                                    FROM отделы INNER JOIN (сотрудники INNER JOIN должности ON
                                                    сотрудники.idДолжности = должности.idДолжности) ON отделы.idОтделы = сотрудники.idОтделы ORDER BY сотрудники.idСотрудники'''
        departments_sql = '''SELECT * FROM отделы'''
        posts_sql = '''SELECT idДолжности, Название, Описание, ТребуемыйОпыт FROM должности'''
        providers_sql = '''SELECT idПоставщики, НаименованиеОрганизации, телефон FROM поставщики'''
        types_sql = '''SELECT idТипы, Наименование, Описание FROM типы'''
        location_sql = '''SELECT idРасположения, Корпус, Этаж, Кабинет FROM расположения'''
        property_sql = '''SELECT имущество.idИмущество, поставщики.НаименованиеОрганизации,
                         concat(сотрудники.Фамилия,' ', сотрудники.Имя,' ', сотрудники.Отчество),
                                     concat(расположения.Корпус, расположения.Этаж, расположения.Кабинет) AS Расположение, типы.Наименование,
                                      имущество.Состояние, имущество.Стоимость 
                                     FROM сотрудники INNER JOIN (типы INNER JOIN (расположения INNER JOIN 
                                    (имущество INNER JOIN поставщики on имущество.idПоставщики = поставщики.idПоставщики)
                                     ON расположения.idРасположения = имущество.idРасположения) ON типы.idТипы = имущество.idТипы)
                                      ON сотрудники.idСотрудники = имущество.idОтветственноеЛицо ORDER BY имущество.idИмущество'''
        queries = {
            "Сотрудники": staff_sql,
            "Отделы": departments_sql,
            "Должности": posts_sql,
            "Имущество": property_sql,
            "Поставщики": providers_sql,
            "Расположения": location_sql,
            "Типы имущества": types_sql,
            "ИКомиссии": commission_sql,
            "Имущество2": property2
        }
        cursor = self.db.cursor()

        for key in queries.keys():
            cursor.execute(queries[key])
            self.tables_with_id[key] = cursor.fetchall()

        staff_sql = '''SELECT idСотрудники, Имя, Фамилия, Отчество, ОпытРаботы, должности.idДолжности, отделы.idОтделы
                                                            FROM отделы INNER JOIN (сотрудники INNER JOIN должности ON
                                                            сотрудники.idДолжности = должности.idДолжности) ON отделы.idОтделы = сотрудники.idОтделы ORDER BY сотрудники.idСотрудники'''
        departments_sql = '''SELECT * FROM отделы'''
        posts_sql = '''SELECT idДолжности, Название, Описание, ТребуемыйОпыт FROM должности'''
        providers_sql = '''SELECT idПоставщики, НаименованиеОрганизации, телефон FROM поставщики'''
        types_sql = '''SELECT idТипы, Наименование, Описание FROM типы'''
        location_sql = '''SELECT idРасположения, Корпус, Этаж, Кабинет FROM расположения'''
        property_sql = '''SELECT имущество.idИмущество, имущество.idПоставщики, имущество.idОтветственноеЛицо,
                                             имущество.idРасположения, имущество.idТипы,
                                              имущество.Состояние, имущество.Стоимость 
                                             FROM сотрудники INNER JOIN (типы INNER JOIN (расположения INNER JOIN 
                                            (имущество INNER JOIN поставщики on имущество.idПоставщики = поставщики.idПоставщики)
                                             ON расположения.idРасположения = имущество.idРасположения) ON типы.idТипы = имущество.idТипы)
                                              ON сотрудники.idСотрудники = имущество.idОтветственноеЛицо ORDER BY имущество.idИмущество'''
        queries = {
            "Сотрудники": staff_sql,
            "Отделы": departments_sql,
            "Должности": posts_sql,
            "Имущество": property_sql,
            "Поставщики": providers_sql,
            "Расположения": location_sql,
            "Типы имущества": types_sql,
            "ИКомиссии": commission_sql,
            "Имущество2": property2
        }
        cursor = self.db.cursor()

        for key in queries.keys():
            cursor.execute(queries[key])
            self.tables_with_all_ids[key] = cursor.fetchall()

    def update_table(self, table_name):
        staff_sql = '''SELECT сотрудники.idСотрудники, Имя, Фамилия, Отчество, ОпытРаботы, должности.Название, отделы.Название,
                       IF (сотрудники.idСотрудники IN (SELECT составкомиссий.idСотрудники FROM составкомиссий),
                        'Да', 'Нет') AS 'Учавствовал в инвентаризации'
                        FROM отделы INNER JOIN (сотрудники INNER JOIN должности ON
                        сотрудники.idДолжности = должности.idДолжности) 
                        ON отделы.idОтделы = сотрудники.idОтделы
                        ORDER BY сотрудники.idСотрудники;'''
        departments_sql = '''SELECT * FROM отделы'''
        posts_sql = '''SELECT Название, Описание, ТребуемыйОпыт FROM должности'''
        providers_sql = '''SELECT НаименованиеОрганизации, телефон FROM поставщики'''
        types_sql = '''SELECT Наименование, Описание FROM типы'''
        location_sql = '''SELECT Корпус, Этаж, Кабинет FROM расположения'''
        property_sql = '''SELECT имущество.idИмущество, поставщики.НаименованиеОрганизации,
                         concat(сотрудники.Фамилия,' ', сотрудники.Имя,' ', сотрудники.Отчество),
                                     concat(расположения.Корпус, расположения.Этаж, расположения.Кабинет) AS Расположение, типы.Наименование,
                                      имущество.Состояние, имущество.Стоимость 
                                     FROM сотрудники INNER JOIN (типы INNER JOIN (расположения INNER JOIN 
                                    (имущество INNER JOIN поставщики on имущество.idПоставщики = поставщики.idПоставщики)
                                     ON расположения.idРасположения = имущество.idРасположения) ON типы.idТипы = имущество.idТипы)
                                      ON сотрудники.idСотрудники = имущество.idОтветственноеЛицо ORDER BY имущество.idИмущество'''
        commission_sql = '''SELECT икомиссии.idИкомиссии, ДатаНачала, ДатаОкончания, ПричинаПроверки, ТипПроверки FROM икомиссии'''
        property2 = '''SELECT * FROM имущество2'''
        queries = {
            "Сотрудники": staff_sql,
            "Отделы": departments_sql,
            "Должности": posts_sql,
            "Имущество": property_sql,
            "Поставщики": providers_sql,
            "Расположения": location_sql,
            "Типы имущества": types_sql,
            "ИКомиссии": commission_sql,
            "Имущество2": property2
        }

        cursor = self.db.cursor()
        cursor.execute(queries[table_name])
        self.tables[table_name] = cursor.fetchall()

        staff_sql = '''SELECT idСотрудники, Имя, Фамилия, Отчество, ОпытРаботы, должности.Название, отделы.Название
                                                            FROM отделы INNER JOIN (сотрудники INNER JOIN должности ON
                                                            сотрудники.idДолжности = должности.idДолжности) ON отделы.idОтделы = сотрудники.idОтделы ORDER BY сотрудники.idСотрудники'''
        departments_sql = '''SELECT * FROM отделы'''
        posts_sql = '''SELECT idДолжности, Название, Описание, ТребуемыйОпыт FROM должности'''
        providers_sql = '''SELECT idПоставщики, НаименованиеОрганизации, телефон FROM поставщики'''
        types_sql = '''SELECT idТипы, Наименование, Описание FROM типы'''
        location_sql = '''SELECT idРасположения, Корпус, Этаж, Кабинет FROM расположения'''
        property_sql = '''SELECT имущество.idИмущество, поставщики.НаименованиеОрганизации,
                                 concat(сотрудники.Фамилия,' ', сотрудники.Имя,' ', сотрудники.Отчество),
                                             concat(расположения.Корпус, расположения.Этаж, расположения.Кабинет) AS Расположение, типы.Наименование,
                                              имущество.Состояние, имущество.Стоимость 
                                             FROM сотрудники INNER JOIN (типы INNER JOIN (расположения INNER JOIN 
                                            (имущество INNER JOIN поставщики on имущество.idПоставщики = поставщики.idПоставщики)
                                             ON расположения.idРасположения = имущество.idРасположения) ON типы.idТипы = имущество.idТипы)
                                              ON сотрудники.idСотрудники = имущество.idОтветственноеЛицо ORDER BY имущество.idИмущество'''
        queries = {
            "Сотрудники": staff_sql,
            "Отделы": departments_sql,
            "Должности": posts_sql,
            "Имущество": property_sql,
            "Поставщики": providers_sql,
            "Расположения": location_sql,
            "Типы имущества": types_sql,
            "ИКомиссии": commission_sql,
            "Имущество2": property2
        }

        cursor.execute(queries[table_name])
        self.tables_with_id[table_name] = cursor.fetchall()

        staff_sql = '''SELECT idСотрудники, Имя, Фамилия, Отчество, ОпытРаботы, должности.idДолжности, отделы.idОтделы
                                                                   FROM отделы INNER JOIN (сотрудники INNER JOIN должности ON
                                                                   сотрудники.idДолжности = должности.idДолжности) ON отделы.idОтделы = сотрудники.idОтделы ORDER BY сотрудники.idСотрудники'''
        departments_sql = '''SELECT * FROM отделы'''
        posts_sql = '''SELECT idДолжности, Название, Описание, ТребуемыйОпыт FROM должности'''
        providers_sql = '''SELECT idПоставщики, НаименованиеОрганизации, телефон FROM поставщики'''
        types_sql = '''SELECT idТипы, Наименование, Описание FROM типы'''
        location_sql = '''SELECT idРасположения, Корпус, Этаж, Кабинет FROM расположения'''
        property_sql = '''SELECT имущество.idИмущество, имущество.idПоставщики, имущество.idОтветственноеЛицо,
                                                    имущество.idРасположения, имущество.idТипы,
                                                     имущество.Состояние, имущество.Стоимость 
                                                    FROM сотрудники INNER JOIN (типы INNER JOIN (расположения INNER JOIN 
                                                   (имущество INNER JOIN поставщики on имущество.idПоставщики = поставщики.idПоставщики)
                                                    ON расположения.idРасположения = имущество.idРасположения) ON типы.idТипы = имущество.idТипы)
                                                     ON сотрудники.idСотрудники = имущество.idОтветственноеЛицо ORDER BY имущество.idИмущество'''
        queries = {
            "Сотрудники": staff_sql,
            "Отделы": departments_sql,
            "Должности": posts_sql,
            "Имущество": property_sql,
            "Поставщики": providers_sql,
            "Расположения": location_sql,
            "Типы имущества": types_sql,
            "ИКомиссии": commission_sql,
            "Имущество2": property2
        }

        cursor.execute(queries[table_name])
        self.tables_with_all_ids[table_name] = cursor.fetchall()

    def get_table(self, table_name, with_id=False, with_all_ids=False):
        if with_id:
            return self.tables_with_id[table_name]
        if with_all_ids:
            return self.tables_with_all_ids[table_name]

        return self.tables[table_name]

    def get_table_name(self, index):
        names = list(self.headers.keys())
        return names[index]

    def insert_into_table(self, table_name, *values, with_id=False):
        cursor = self.db.cursor()
        if table_name == "Типы":
            table = self.get_table("Типы имущества", with_id=True)
        else:
            table = self.get_table(table_name, with_id=True)

        auto_increment = 0
        if table_name != "Имущество" and table_name != "Отделы":
            auto_increment = table[-1][0] + 1

            for i in range(0, len(table) - 1):
                if abs(table[i][0] - table[i + 1][0]) >= 2:
                    auto_increment = table[i][0] + 1
                    break

        cursor.execute(f"SELECT * FROM {table_name}")
        headers = cursor.description

        table_name = table_name[0].lower() + table_name[1:]
        sql = f"INSERT INTO {table_name} ("
        for header in headers:
            sql += header[0] + ", "
        sql = sql[0:-2]
        sql += ") VALUES("
        if with_id:
            for value in values:
                try:
                    sql += f"\'{int(value)}\', "
                except ValueError:
                    sql += f"\'{value}\', "
        else:
            sql += f"\'{int(auto_increment)}\', "
            for value in values:
                try:
                    sql += f"\'{int(value)}\', "
                except ValueError:
                    sql += f"\'{value}\', "
        sql = sql[0:-2]
        sql += ");"
        cursor.execute(sql)
        self.db.commit()
        if table_name == "типы":
            table_name = "Типы имущества"
        table_name = table_name[0].upper() + table_name[1:]
        self.update_table(table_name)

    def get_current_table(self):
        result_table = []
        for i in range(0, self.table.rowCount()):
            record = []
            for j in range(0, self.table.columnCount()):
                value = self.table.item(i, j).text()
                record.append(value)
            result_table.append(record)
        return result_table

    def delete_record_from_table(self, identifier, table_name):
        sql_query = f"DELETE FROM {table_name} WHERE (id{table_name} = {identifier});"
        cursor = self.db.cursor()
        cursor.execute(sql_query)
        self.db.commit()
        if table_name == "Типы":
            table_name = "Типы имущества"
        self.update_table(table_name)

    def update_record_in_table(self, table_name, identifier, row_index, new_value):
        cursor = self.db.cursor()
        cursor.execute(f"SELECT * FROM {table_name}")
        headers = []
        for header in cursor.description:
            headers.append(header[0])
        if isinstance(new_value, str):
            sql_query = f"UPDATE {table_name} SET {headers[row_index]} = '{new_value}' WHERE {headers[0]} = {identifier};"
        else:
            sql_query = f"UPDATE {table_name} SET {headers[row_index]} = {new_value} WHERE {headers[0]} = {identifier};"
        cursor.execute(sql_query)
        self.db.commit()
        if table_name == "Типы":
            table_name = "Типы имущества"
        self.update_table(table_name)

    def is_id_in_table(self, table_name, identifier):
        table = self.get_table(table_name, with_id=True)
        existing_ids = [record[0] for record in table]
        if identifier in existing_ids:
            return True
        return False

    def disable_change_some_cells(self, **kwargs):
        if not isinstance(kwargs['rows'], list):
            if kwargs['last_row'] is None:
                kwargs['last_row'] = self.table.rowCount()
            kwargs['rows'] = range(kwargs['first_row'], kwargs['last_row'])

        if not isinstance(kwargs['columns'], list):
            if kwargs['last_column'] is None:
                kwargs['last_column'] = self.table.columnCount()
            kwargs['columns'] = range(kwargs['first_column'], kwargs['last_column'])

        for i in kwargs['rows']:
            for j in kwargs['columns']:
                self.table.item(i, j).setFlags(
                    self.table.item(i, j).flags() & ~QtCore.Qt.ItemIsEditable)

    @staticmethod
    def left_found_records(table, searching_word, row_number):
        result_table = []
        for record in table:
            k = 0
            index = 0
            for letter in searching_word:
                for i in range(index, len(record[row_number - 1])):
                    if letter.lower() == record[row_number - 1][i].lower():
                        k += 1
                        index = i + 1
                        break
            if k == len(searching_word):
                result_table.append(record)

        return result_table

    @staticmethod
    def convert_table_to_string(table):
        new_table = []
        for record in table:
            new_record = []
            for value in record:
                new_record.append(str(value))
            new_table.append(new_record)

        return new_table

    @staticmethod
    def descending_sort(table, row_number):
        for i in range(1, len(table)):
            f = 0
            for j in range(0, len(table) - i):
                if table[j][row_number - 1] < table[j + 1][row_number - 1]:
                    table[j], table[j + 1] = table[j + 1], table[j]
                    f = 1
            if f == 0:
                break

        return table

    @staticmethod
    def ascending_sort(table, row_number):
        for i in range(1, len(table)):
            f = 0
            for j in range(0, len(table) - i):
                if table[j][row_number - 1] > table[j + 1][row_number - 1]:
                    table[j], table[j + 1] = table[j + 1], table[j]
                    f = 1
            if f == 0:
                break

        return table

    @staticmethod
    def left_only_from_to_records(bottom_value, top_value, table, row_number):
        result_table = []
        for record in table:
            if bottom_value <= record[row_number - 1] <= top_value:
                result_table.append(record)

        return result_table

    @staticmethod
    def left_only_one_value(table, value, row_number):
        result_table = []
        for record in table:
            if record[row_number - 1] == value:
                result_table.append(record)

        return result_table

    def color_by_matching_with_another_table(self, table, color):
        r = color[0]
        g = color[1]
        b = color[2]
        current_table = self.get_current_table()
        for i in range(0, len(current_table)):
            for j in range(0, len(current_table[i])):
                if isinstance(current_table[i][j], str) and j == 1:
                    temp = current_table[i][j]
                    temp = temp.replace('(', '')
                    temp = temp.replace(')', '')
                    temp = temp.strip('+')
                    temp = temp.replace('-', '')
                    current_table[i][j] = temp
        for i, record in enumerate(current_table):
            for rec in table:
                k = 0
                for q in range(0, len(rec)):
                    if rec[q] == record[q]:
                        k += 1
                if k == len(rec):
                    for j in range(0, self.table.columnCount()):
                        self.table.item(i, j).setBackground(QtGui.QColor(r, g, b))
                    break

    def display_table(self, table_name, table=None):
        needed_columns = {
            "Сотрудники": [None, 0, None, [0, 7], 0, 0],
            "Должности": [None, 0, None, None, 2, None],
            "Отделы": [None, 0, None, None, 0, 1],
            "Имущество": [None, 0, None, [0, 4], 0, None],
            "Расположения": [None, 0, None, None, 0, None],
            "Типы имущества": [None, 0, None, None, 0, 1],
            "ИКомиссии": [None, 0, None, None, 0, None]
        }
        if table is None:
            table = self.get_table(table_name)
        self.table.setRowCount(0)
        self.table.setColumnCount(0)
        if table_name == "Поставщики":
            for row_number, row_data in enumerate(table):
                self.table.insertRow(row_number)
                for column_number, data in enumerate(row_data):
                    if self.table.columnCount() <= column_number:
                        self.table.setColumnCount(self.table.columnCount() + 1)
                    if column_number == 1:
                        data = '+{}({})-{}-{}-{}'.format(data[:3], data[3:5], data[5:8], data[8:10], data[10:12])
                    self.table.setItem(row_number, column_number, QTableWidgetItem(str(data)))
        else:
            for row_number, row_data in enumerate(table):
                self.table.insertRow(row_number)
                for column_number, data in enumerate(row_data):
                    if self.table.columnCount() <= column_number:
                        self.table.setColumnCount(self.table.columnCount() + 1)
                    if data is None:
                        data = ""
                    self.table.setItem(row_number, column_number, QTableWidgetItem(str(data)))
        #
        self.table.setHorizontalHeaderLabels(self.headers[table_name])
        self.table.resizeColumnsToContents()
        if table_name in needed_columns.keys():
            temp = needed_columns[table_name]
            self.disable_change_some_cells(rows=temp[0], first_row=temp[1], last_row=temp[2],
                                           columns=temp[3], first_column=temp[4], last_column=temp[5])


class DateDialog(QtWidgets.QDialog):
    def __init__(self, parent=None):
        super(DateDialog, self).__init__(parent)

        layout = QtWidgets.QVBoxLayout(self)

        font = QtGui.QFont()
        font.setFamily("Comic Sans MS")
        font.setPointSize(9)
        font.setBold(True)
        font.setWeight(75)

        self.label1 = QtWidgets.QLabel(self)
        self.label1.setText("Дата начала")
        self.label1.setFont(font)
        layout.addWidget(self.label1)

        self.start_date = QtWidgets.QDateEdit(self)
        self.start_date.setDate(datetime.datetime.now())
        layout.addWidget(self.start_date)

        self.label2 = QtWidgets.QLabel(self)
        self.label2.setText("Дата окончания")
        self.label2.setFont(font)
        layout.addWidget(self.label2)

        self.end_date = QtWidgets.QDateEdit(self)
        self.end_date.setDate(datetime.datetime.now())
        layout.addWidget(self.end_date)

        buttons = QtWidgets.QDialogButtonBox(
            QtWidgets.QDialogButtonBox.Ok | QtWidgets.QDialogButtonBox.Cancel,
            QtCore.Qt.Horizontal, self)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)

    def date1(self):
        return self.start_date.text(), self.end_date.text()

    @staticmethod
    def get_date(parent=None):
        dialog = DateDialog(parent)
        result = dialog.exec_()
        start_date, end_date = dialog.date1()
        return start_date, end_date, result == QtWidgets.QDialog.Accepted


#
# QTabBar::tab:selected  {
#     background-color:rgb(255, 255, 0);
# }


qss = '''
* {
    color: rgb(0, 0, 0);
    background-color: rgb(190, 190, 190);
}
QPushButton,
QLineEdit,
QComboBox {
    background-color: rgb(150, 150, 150);
}
QTableWidget {
    background-color: rgb(240, 240, 240);
}
QTabBar::tab {
    background-color:rgb(140, 140, 140);
    border-right: 1px solid black;
    border-top: 1px solid black;
    border-bottom: 1px solid black;
    padding: 5px 10px;
}
QHeaderView::section {
    border-style: solid;
    background-color: rgb(190, 190, 190);
    border-bottom: 1px solid black;
    border-right: 1px solid black;

 }

QTabBar::tab::selected {
    background-color:qlineargradient( x1: 0, y1: 0, x2: 0, y2: 1,
                                        stop: 0 rgb(190, 190, 190), stop: 1 rgb(100, 100, 100));
}
QTabBar::tab::disabled {
    background-color:rgb(250, 250, 250);
}
QLineEdit {
    border: 1px solid black;
}
'''

if __name__ == "__main__":
    previous_cell = []
    app = QtWidgets.QApplication(sys.argv)
    app.setStyleSheet(qss)
    MainWindow = QtWidgets.QMainWindow()
    # ui = Ui_MainWindow()
    # ui.setupUi(MainWindow)
    # application = Program()
    # application.show()
    a = Program()
    sys.exit(app.exec_())

# app = QtWidgets.QApplication([])

# sys.exit(app.exec())
